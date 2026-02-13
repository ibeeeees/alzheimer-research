"""Generate a 2-page .docx research proposal."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# -- Page margins --
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15

# -- Title --
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Ordinal Multi-Modal Deep Learning for Alzheimer\'s Disease Severity Prediction')
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'
title.paragraph_format.space_after = Pt(2)

# -- Authors --
authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = authors.add_run('Ibe Mohammed Ali, Kubra Sag, Poorav Rawat')
run.font.size = Pt(11)
run.font.name = 'Times New Roman'
authors.paragraph_format.space_after = Pt(10)


def add_heading_text(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Inches(0.3)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
    return p


# -- Project Type --
add_heading_text('1. Project Type')
add_body(
    'This is a research-flavor project. We develop a novel ordinal multi-task deep learning '
    'framework for Alzheimer\'s disease severity prediction using multi-modal biomedical data. '
    'The core methods include CORAL ordinal regression, homoscedastic multi-task uncertainty '
    'weighting, and discrete-time survival analysis. We extend these methods into a unified '
    'cross-cohort system that has not been explored in the Alzheimer\'s disease literature.'
)

# -- ML Problem --
add_heading_text('2. Machine Learning Problem')
add_body(
    'The problem is formulated as supervised ordinal regression with auxiliary survival prediction. '
    'Given a structural MRI brain scan or a speech recording from a cognitive assessment, the model '
    'must predict an ordered Alzheimer\'s disease severity stage: CDR 0 (cognitively normal), '
    'CDR 0.5 (very mild), CDR 1.0 (mild dementia), or CDR 2.0 and above (moderate to severe). '
    'For patients at the MCI stage (CDR 0.5), the model additionally predicts a time-varying '
    'probability of conversion to Alzheimer\'s dementia over a 36-month horizon. Unlike standard '
    'multi-class classification with softmax, the ordinal formulation explicitly penalizes distant '
    'misclassifications more heavily than adjacent ones, respecting the progressive nature of the '
    'disease.'
)

# -- Goals and Motivation --
add_heading_text('3. Goals and Motivation')
add_body(
    'Convolutional neural networks are widely used for MRI-based Alzheimer\'s diagnosis, but most '
    'existing approaches treat disease stages as independent categories using softmax classification. '
    'This is problematic for several reasons. First, misclassifying a cognitively '
    'normal patient as having moderate dementia is clinically far worse than misclassifying them as '
    'very mild, yet cross-entropy loss treats both errors equally. Second, single-task binary models '
    'that classify Alzheimer\'s disease versus cognitively normal discard the ordinal structure entirely '
    'and cannot model disease progression. Third, most multimodal approaches require paired data from '
    'the same subjects, discarding the majority of available unimodal data.'
)
add_body(
    'Our system addresses these challenges by using CORAL ordinal regression to preserve severity '
    'ordering with learned thresholds that map a continuous severity score to ordered stages; jointly '
    'training a discrete-time survival head for MCI-to-AD conversion prediction that handles '
    'right-censored subjects natively; integrating structural MRI and speech biomarkers from separate '
    'cohorts without requiring paired subjects, using cross-cohort alignment via shared ordinal heads '
    'and class-conditioned Maximum Mean Discrepancy; and balancing multiple '
    'task losses automatically through learned homoscedastic uncertainty weights. No prior published '
    'work combines all four of these elements in a single framework.'
)

# -- Methodology --
add_heading_text('4. Methodology and Models')
add_body(
    'Structural MRI volumes, resampled to 128 by 128 by 128 voxels, are processed by a 3D ResNet-18 '
    'adapted from the video recognition domain to single-channel medical imaging. '
    'The pretrained Kinetics-400 weights are transferred by averaging the three RGB input channels into '
    'a single grayscale channel. The backbone produces a 256-dimensional embedding per scan. For patients '
    'with multiple longitudinal visits, per-visit embeddings are aggregated by a GRU recurrent network '
    'with sinusoidal time encoding that captures disease trajectory across irregular visit intervals, '
    'an approach that handles the variable timing between clinical visits without requiring fixed-step '
    'assumptions.'
)
add_body(
    'Speech recordings from a Cookie Theft picture description task are represented as a concatenation '
    'of four feature streams: wav2vec 2.0 learned acoustic embeddings of 768 dimensions, '
    'Sentence-BERT transcript embeddings of 384 dimensions, handcrafted acoustic features of 216 '
    'dimensions including MFCCs, prosody, voice quality, and temporal measures, and handcrafted linguistic '
    'features of 14 dimensions including lexical diversity, syntactic complexity, semantic coherence, and '
    'fluency markers. This 1382-dimensional combined vector is projected to 256 '
    'dimensions through a two-layer MLP with layer normalization and GELU activation.'
)
add_body(
    'Both encoders feed into a shared CORAL ordinal head that predicts a single continuous severity score, '
    'which three learnable thresholds partition into four ordered CDR stages. The CORAL loss computes '
    'binary cross-entropy at each threshold, ensuring rank consistency across cumulative probabilities. '
    'For MCI patients, a separate survival head predicts interval-specific hazard probabilities over six '
    '6-month intervals spanning a 36-month window, handling right-censored subjects without imputation. '
    'Since MRI and speech data come from non-overlapping populations, the shared CORAL head implicitly '
    'aligns representations by forcing both modalities onto the same severity scale, while class-conditioned '
    'MMD explicitly minimizes distributional distance between same-severity embeddings from different '
    'modalities. All task losses are combined via homoscedastic uncertainty weighting, where learned '
    'variance parameters automatically downweight noisier tasks during training.'
)

# -- Schedule --
add_heading_text('5. Schedule')
add_body(
    'The project follows an eleven-week schedule. Weeks 1 through 2 focus on data acquisition, MRI '
    'preprocessing, and speech feature extraction. Weeks 3 through 4 involve training unimodal baselines '
    'including a standard softmax CNN, an MRI ordinal model, and a speech MLP. Weeks 5 through 6 are '
    'dedicated to implementing CORAL ordinal regression and comparing it against the softmax baseline. '
    'Weeks 7 through 8 cover building the multi-task model with ordinal and survival heads and '
    'implementing cross-cohort MMD alignment. Weeks 9 through 10 focus on full system training, '
    'hyperparameter tuning, and ablation studies. Week 11 is reserved for final evaluation, statistical '
    'analysis, and report preparation.'
)

# -- Datasets --
add_heading_text('6. Datasets')
add_body(
    'The primary imaging dataset is SCAN from the National Alzheimer\'s Coordinating Center, which '
    'contains approximately 29,000 3D T1-weighted MRI scans from roughly 10,000 subjects with '
    'longitudinal follow-up. Each scan is linked to the NACC Uniform Data Set '
    'providing CDR scores, diagnosis codes, and neuropsychological assessments. SCAN\'s centralized '
    'acquisition and quality control pipeline eliminates the site-harmonization confounds present in '
    'ADNI-based studies. The speech dataset is the DementiaBank Pitt Corpus, which contains approximately '
    '550 Cookie Theft picture description recordings from around 270 subjects, including 170 with '
    'probable Alzheimer\'s disease and 100 healthy controls. We will additionally benchmark on the Kaggle '
    'Alzheimer\'s Multiclass MRI Dataset containing approximately 44,000 augmented 2D slices across four '
    'severity classes. Patient-level splits are enforced across all datasets to prevent data leakage, '
    'with stratification on baseline CDR and conversion status.'
)

# -- Evaluation --
add_heading_text('7. Evaluation')
add_body(
    'Model performance will be evaluated using Mean Absolute Error to measure ordinal stage distance, '
    'Quadratic Weighted Kappa to assess agreement while penalizing large ordinal errors, and '
    'classification accuracy as a standard reference metric. For the survival component, we use '
    'Harrell\'s concordance index for ranking accuracy and time-dependent AUC at 12, 24, and 36 months '
    'for conversion prediction at clinically relevant horizons. Baseline comparisons include a standard '
    'multi-class CNN with softmax loss, unimodal MRI-only and speech-only models, multi-task training '
    'without ordinal constraints, and multi-modal training without cross-cohort alignment. We will also '
    'vary the training set size at 10, 25, 50, 75, and 100 percent to analyze data efficiency and '
    'learning curves.'
)

# -- Resources --
add_heading_text('8. Resources')
add_body(
    'Training will be performed on Google Colab Pro using a T4 GPU with 16 GB of VRAM. Phase 1 MRI '
    'pretraining requires approximately 9 hours in a single session, and Phase 2 multi-task training '
    'requires approximately 50 hours across 3 to 4 sessions using checkpoint-resume. The software stack '
    'includes Python, PyTorch, torchvision, scikit-learn, nibabel for neuroimaging I/O, Parselmouth for '
    'acoustic analysis, and the Hugging Face transformers library for wav2vec 2.0 and Sentence-BERT. '
    'Version control is managed through GitHub.'
)

# -- Workload --
add_heading_text('9. Workload Distribution')
add_body(
    'Ibe Mohammed Ali is responsible for the MRI encoder, 3D ResNet-18 adaptation, longitudinal temporal '
    'module, cross-cohort alignment via MMD, and overall system integration. Kubra Sag is responsible '
    'for the speech encoder, acoustic and linguistic feature extraction, wav2vec 2.0 and Sentence-BERT '
    'pipelines, and DementiaBank preprocessing. Poorav Rawat is responsible for the ordinal regression '
    'implementation using CORAL, the survival head, multi-task loss balancing, evaluation metrics, and '
    'baseline comparisons. All members collaborate on experimental design, ablation studies, and final '
    'report writing.'
)

# -- References --
add_heading_text('References')
refs = [
    'Yee, E., Ma, D., Popuri, K., et al. (2025). DiaMond: Dementia diagnosis with multi-modal vision Transformers using MRI and PET. IEEE Transactions on Medical Imaging, 44(1). This recent work represents the state of the art in multi-modal deep learning for dementia diagnosis, using paired MRI and PET with Vision Transformers. Our approach differs fundamentally by operating across non-overlapping cohorts, using ordinal regression instead of categorical classification, and incorporating speech biomarkers and survival modeling.',
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

output_path = '/Users/ibe/antigravity/alzheimer-research/PROPOSAL.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
